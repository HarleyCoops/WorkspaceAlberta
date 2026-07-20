"""E2B-backed bid room processing for isolated tender package work.

The "bid room" is the heaviest tool in WorkspaceAlberta: it opens a tender's
attachment package (PDFs, DOCX, XLSX, ZIPs), extracts compliance-relevant
text, and runs a structured Cohere Command A+ review — all inside a
short-lived E2B sandbox so unknown user-facing files never touch the
always-on MCP service.

How it works, end to end:

1.  **Payload build (host side).** :func:`build_canadabuys_bid_room_payload`
    or :func:`build_apc_bid_room_payload` turns a cached CanadaBuys row or a
    live APC detail response plus the saved business profile into a JSON
    payload: opportunity metadata, attachment URLs (capped at
    ``MAX_ATTACHMENTS``), profile context, and processing limits.
2.  **Sandbox execution.** :func:`run_live_bid_room_process` boots an E2B
    sandbox, injects ``SANDBOX_PROCESSOR`` (a self-contained Python script,
    stored as a raw string in this module) with the payload substituted in,
    and runs it under a command timeout. The processor downloads attachments,
    extracts text (pdfminer/python-docx/openpyxl installed on demand via
    :func:`ensure_package`), walks ZIPs up to ``MAX_ZIP_MEMBERS``, and builds
    an evidence bundle of normalized document text.
3.  **In-sandbox Cohere review.** ``call_cohere`` (inside the processor) calls
    Command A+ with read-only evidence tools (``search_extracted_documents``,
    ``get_bid_evidence``) and a strict JSON schema
    (``COHERE_RESPONSE_SCHEMA``). The model must ground its answer in
    extracted text; responses are validated by ``validate_cohere_analysis``
    against ``REQUIRED_COHERE_FIELDS`` (bid recommendation, fit score,
    requirements, risks, missing info, deadlines, questions, next actions).
4.  **Artifact return.** The processor prints a single JSON artifact to
    stdout; :func:`parse_artifact` recovers it, ``validate_bid_room_artifact``
    checks its shape, and :func:`render_bid_room_markdown` formats the
    :class:`BidRoomSandboxResult` for chat display. The sandbox is killed
    unless ``keep_alive`` was requested.

Safety limits: ``MAX_FILE_BYTES`` (25 MB per file), ``MAX_COHERE_CHARS``
(80k prompt chars), ``MAX_ATTACHMENTS`` (5), and per-command timeouts. All
are enforced inside the sandbox as well as on the host.

Requires ``E2B_API_KEY``; Cohere analysis inside the sandbox additionally
requires ``COHERE_API_KEY`` (see :func:`has_e2b_api_key` /
:func:`has_cohere_api_key`). Without a Cohere key the sandbox still extracts
and returns document evidence — only the model review is skipped.
"""

from __future__ import annotations

import json
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib.parse import unquote, urlparse

ROOT_DIR = Path(__file__).resolve().parents[1]

COHERE_MODEL = "command-a-plus-05-2026"
COHERE_CHAT_URL = "https://api.cohere.com/v2/chat"
MAX_ATTACHMENTS = 5
MAX_FILE_BYTES = 25 * 1024 * 1024
MAX_COHERE_CHARS = 80_000
REQUIRED_COHERE_FIELDS = (
    "bid_recommendation",
    "fit_score",
    "requirements",
    "risks",
    "missing_information",
    "deadlines",
    "questions_to_ask",
    "next_actions",
)
COHERE_RESPONSE_SCHEMA = {
    "type": "object",
    "properties": {
        "bid_recommendation": {"type": "string"},
        "fit_score": {"type": "integer"},
        "requirements": {"type": "array", "items": {"type": "string"}},
        "risks": {"type": "array", "items": {"type": "string"}},
        "missing_information": {"type": "array", "items": {"type": "string"}},
        "deadlines": {"type": "array", "items": {"type": "string"}},
        "questions_to_ask": {"type": "array", "items": {"type": "string"}},
        "next_actions": {"type": "array", "items": {"type": "string"}},
    },
    "required": list(REQUIRED_COHERE_FIELDS),
    "additionalProperties": False,
}
COHERE_RESPONSE_FORMAT = {
    "type": "json_object",
    "schema": COHERE_RESPONSE_SCHEMA,
}


SAMPLE_DOCUMENTS = [
    {
        "name": "opportunity-notice.txt",
        "text": (
            "Alberta public sector opportunity for structural steel fabrication. "
            "Closing date: 2026-06-18 14:00 MT. Mandatory site meeting: "
            "2026-05-29 10:00 MT. The supplier shall provide shop drawings, "
            "CSA W47.1 certification, commercial general liability insurance, "
            "and proof of bonding capacity."
        ),
    },
    {
        "name": "submission-instructions.txt",
        "text": (
            "Responses must include a signed bid form, safety program summary, "
            "three comparable project references, delivery schedule, and price "
            "breakdown. Late submissions will not be accepted. Questions must be "
            "submitted five business days before closing."
        ),
    },
    {
        "name": "company-profile.txt",
        "text": (
            "Company profile: Edmonton steel fabrication shop specializing in "
            "structural beams, railings, custom metalwork, commercial construction, "
            "and shop drawings. Located in Edmonton, Alberta."
        ),
    },
]


SANDBOX_PROCESSOR = r"""
import hashlib
import html
import json
import os
import re
import subprocess
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import unquote, urlparse
from urllib.request import Request, urlopen

payload = json.loads(__PAYLOAD_JSON__)
work_dir = Path("/tmp/workspacealberta-bid-room")
download_dir = work_dir / "downloads"
extract_dir = work_dir / "extract"
download_dir.mkdir(parents=True, exist_ok=True)
extract_dir.mkdir(parents=True, exist_ok=True)

MAX_FILE_BYTES = int(payload.get("limits", {}).get("max_file_bytes", 25 * 1024 * 1024))
MAX_COHERE_CHARS = int(payload.get("limits", {}).get("max_cohere_chars", 80000))
MAX_ZIP_MEMBERS = 12
MAX_DOC_CHARS = 24000
USER_AGENT = "WorkspaceAlberta-BidRoom/0.1"
REQUIRED_COHERE_FIELDS = (
    "bid_recommendation",
    "fit_score",
    "requirements",
    "risks",
    "missing_information",
    "deadlines",
    "questions_to_ask",
    "next_actions",
)


def ensure_package(import_name, package_name=None):
    try:
        __import__(import_name)
        return
    except ImportError:
        pass
    subprocess.run(
        [sys.executable, "-m", "pip", "install", "-q", package_name or import_name],
        check=True,
        timeout=120,
    )


def safe_name(value, fallback="document"):
    parsed = urlparse(value or "")
    name = unquote(Path(parsed.path).name) if parsed.path else ""
    name = name or fallback
    name = re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("._")
    return name[:120] or fallback


def sha256_bytes(data):
    return hashlib.sha256(data).hexdigest()


def read_url(url):
    request = Request(url, headers={"User-Agent": USER_AGENT, "Accept": "*/*"})
    try:
        with urlopen(request, timeout=90) as response:
            content_length = response.headers.get("Content-Length")
            if content_length and int(content_length) > MAX_FILE_BYTES:
                return None, response.headers, f"file too large from content-length: {content_length}"
            data = response.read(MAX_FILE_BYTES + 1)
            if len(data) > MAX_FILE_BYTES:
                return None, response.headers, f"file exceeded {MAX_FILE_BYTES} byte limit"
            return data, response.headers, ""
    except HTTPError as exc:
        return None, {}, f"HTTP {exc.code}"
    except URLError as exc:
        return None, {}, f"URL error: {exc.reason}"
    except Exception as exc:
        return None, {}, str(exc)


def decode_text(data):
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def html_to_text(raw):
    raw = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw)
    raw = re.sub(r"(?is)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?is)</p\s*>", "\n", raw)
    raw = re.sub(r"(?is)<.*?>", " ", raw)
    raw = html.unescape(raw)
    return "\n".join(line.strip() for line in raw.splitlines() if line.strip())


def extract_pdf(path):
    ensure_package("cryptography")
    ensure_package("pypdf")
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages[:80]:
        pages.append(page.extract_text() or "")
        if sum(len(item) for item in pages) >= MAX_DOC_CHARS:
            break
    return "\n".join(pages)


def extract_docx(path):
    ensure_package("docx", "python-docx")
    import docx

    doc = docx.Document(str(path))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables[:20]:
        for row in table.rows[:100]:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx(path):
    ensure_package("openpyxl")
    import openpyxl

    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    rows = []
    for sheet in wb.worksheets[:8]:
        rows.append(f"Sheet: {sheet.title}")
        for index, row in enumerate(sheet.iter_rows(values_only=True), 1):
            if index > 250:
                break
            cells = [str(cell) for cell in row if cell not in (None, "")]
            if cells:
                rows.append(" | ".join(cells))
        if sum(len(item) for item in rows) >= MAX_DOC_CHARS:
            break
    return "\n".join(rows)


def extract_from_path(path, content_type=""):
    suffix = path.suffix.lower()
    if suffix == ".pdf" or "pdf" in content_type:
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".xlsx":
        return extract_xlsx(path)
    if suffix in {".html", ".htm"} or "html" in content_type:
        return html_to_text(path.read_text(encoding="utf-8", errors="replace"))
    if suffix in {".txt", ".md", ".csv", ".json", ".xml"}:
        return path.read_text(encoding="utf-8", errors="replace")
    data = path.read_bytes()
    if b"\x00" not in data[:2048]:
        return decode_text(data)
    raise ValueError(f"unsupported file type: {suffix or content_type or 'unknown'}")


def extract_zip(path):
    outputs = []
    with zipfile.ZipFile(path) as archive:
        for member in archive.infolist()[:MAX_ZIP_MEMBERS]:
            if member.is_dir() or member.file_size > MAX_FILE_BYTES:
                continue
            member_name = safe_name(member.filename, "zip_member")
            target = extract_dir / f"{path.stem}_{member_name}"
            with archive.open(member) as source:
                target.write_bytes(source.read(MAX_FILE_BYTES + 1))
            try:
                text = extract_from_path(target)
                if text.strip():
                    outputs.append(f"--- {member.filename} ---\n{text[:MAX_DOC_CHARS]}")
            except Exception as exc:
                outputs.append(f"--- {member.filename} ---\n[Extraction failed: {exc}]")
    return "\n\n".join(outputs)


def extract_document_text(path, content_type=""):
    if path.suffix.lower() == ".zip" or "zip" in content_type:
        return extract_zip(path)
    return extract_from_path(path, content_type)


def normalize_line(text):
    return " ".join(str(text).split())


def extract_evidence(text_documents, profile):
    requirement_patterns = [
        r"\bmust\b[^.\n]{0,360}[.\n]",
        r"\bshall\b[^.\n]{0,360}[.\n]",
        r"\bmandatory\b[^.\n]{0,360}[.\n]",
        r"\brequired\b[^.\n]{0,360}[.\n]",
        r"\binsurance\b[^.\n]{0,360}[.\n]",
        r"\bbond(?:ing)?\b[^.\n]{0,360}[.\n]",
        r"\bcertification\b[^.\n]{0,360}[.\n]",
        r"\bsafety\b[^.\n]{0,360}[.\n]",
        r"\bsecurity clearance\b[^.\n]{0,360}[.\n]",
    ]
    deadline_patterns = [
        r"(?:closing date|closing|site meeting|questions? must be submitted|deadline|due date)[: ]+[^.\n]{0,240}[.\n]?",
        r"\b20\d{2}-\d{2}-\d{2}(?:[T ][0-9:]{4,8})?",
    ]
    requirements = []
    deadlines = []
    matched_terms = set()
    keywords = [str(word).lower() for word in profile.get("keywords", []) if str(word).strip()]
    text_for_model = []

    for document in text_documents:
        text = document.get("text", "")
        lowered = text.lower()
        for keyword in keywords:
            if keyword in lowered:
                matched_terms.add(keyword)
        for pattern in requirement_patterns:
            for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                line = normalize_line(match.group(0))
                if line:
                    requirements.append({"source": document["name"], "text": line[:500]})
        for pattern in deadline_patterns:
            for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                line = normalize_line(match.group(0))
                if line:
                    deadlines.append({"source": document["name"], "text": line[:400]})
        text_for_model.append(f"## {document['name']}\n{text[:MAX_DOC_CHARS]}")

    def dedupe(items):
        seen = set()
        output = []
        for item in items:
            key = (item["source"], item["text"].lower())
            if key in seen:
                continue
            seen.add(key)
            output.append(item)
        return output

    return {
        "requirements": dedupe(requirements)[:80],
        "deadlines": dedupe(deadlines)[:40],
        "matched_terms": sorted(matched_terms),
        "text_for_model": "\n\n".join(text_for_model)[:MAX_COHERE_CHARS],
    }


def parse_model_json(content):
    content = content.strip()
    content = re.sub(r"<\|START_THINKING\|>.*?<\|END_THINKING\|>", "", content, flags=re.DOTALL)
    content = re.sub(r"<START_THINKING>.*?<END_THINKING>", "", content, flags=re.DOTALL)
    content = re.sub(r"</?[^>]*RESPONSE[^>]*>", "", content)
    content = content.strip()
    content = re.sub(r"^```(?:json)?\s*", "", content)
    content = re.sub(r"\s*```$", "", content)
    try:
        return json.loads(content)
    except json.JSONDecodeError:
        start = content.find("{")
        end = content.rfind("}")
        if start == -1 or end == -1 or end <= start:
            raise ValueError(f"Cohere did not return parseable JSON. Snippet: {content[:400]!r}")
        candidate = content[start:end + 1]
        try:
            return json.loads(candidate)
        except json.JSONDecodeError as exc:
            raise ValueError(
                f"Cohere returned an invalid JSON object: {exc}. Snippet: {candidate[:800]!r}"
            ) from exc


def validate_cohere_analysis(data):
    if not isinstance(data, dict):
        raise ValueError("Cohere analysis must be a JSON object.")
    missing = [field for field in REQUIRED_COHERE_FIELDS if field not in data]
    if missing:
        raise ValueError(f"Cohere analysis missing fields: {', '.join(missing)}")
    if not isinstance(data["bid_recommendation"], str):
        raise ValueError("bid_recommendation must be a string.")
    try:
        data["fit_score"] = int(data["fit_score"])
    except (TypeError, ValueError) as exc:
        raise ValueError("fit_score must be an integer.") from exc
    data["fit_score"] = max(0, min(100, data["fit_score"]))
    for field in (
        "requirements",
        "risks",
        "missing_information",
        "deadlines",
        "questions_to_ask",
        "next_actions",
    ):
        if isinstance(data[field], str):
            data[field] = [data[field]]
        if not isinstance(data[field], list):
            raise ValueError(f"{field} must be a list.")
    return data


COHERE_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "search_extracted_documents",
            "description": (
                "Search the extracted tender documents for source-grounded snippets. "
                "Use this before making bid/no-bid claims."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Terms to search for in the extracted tender text.",
                    },
                    "top_k": {
                        "type": "integer",
                        "description": "Maximum number of matching snippets to return.",
                    },
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "get_bid_evidence",
            "description": (
                "Return deterministic evidence extracted before the model call, "
                "including requirements, deadlines, matched profile terms, and document summaries."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "section": {
                        "type": "string",
                        "enum": [
                            "requirements",
                            "deadlines",
                            "matched_terms",
                            "documents",
                            "opportunity",
                            "profile",
                        ],
                        "description": "Evidence section to return.",
                    },
                    "top_k": {
                        "type": "integer",
                        "description": "Maximum number of records to return for list sections.",
                    },
                },
                "required": ["section"],
            },
        },
    },
]


def cohere_post(api_key, request_payload):
    request = Request(
        payload.get("cohere", {}).get("endpoint", "https://api.cohere.com/v2/chat"),
        data=json.dumps(request_payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "User-Agent": USER_AGENT,
        },
        method="POST",
    )
    try:
        with urlopen(request, timeout=180) as response:
            return json.loads(response.read().decode("utf-8"))
    except HTTPError as exc:
        body = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"Cohere returned HTTP {exc.code}: {body[:500]}") from exc


def cohere_post_with_fallbacks(api_key, request_payload):
    try:
        return cohere_post(api_key, request_payload)
    except RuntimeError as exc:
        message = str(exc)
        if "tool_choice is not supported" in message and "tool_choice" in request_payload:
            retry_payload = dict(request_payload)
            retry_payload.pop("tool_choice", None)
            return cohere_post_with_fallbacks(api_key, retry_payload)
        if "strict_tools" in message and "strict_tools" in request_payload:
            retry_payload = dict(request_payload)
            retry_payload.pop("strict_tools", None)
            return cohere_post_with_fallbacks(api_key, retry_payload)
        raise


def extract_message_text(message):
    content = message.get("content", "")
    if isinstance(content, list):
        text_parts = []
        for part in content:
            if isinstance(part, dict):
                if part.get("type") == "text" and "text" in part:
                    text_parts.append(str(part["text"]))
                elif "text" in part and part.get("type") != "thinking":
                    text_parts.append(str(part["text"]))
                continue
            text_parts.append(str(part))
        return "".join(text_parts)
    return str(content or "")


def normalize_tool_call(tool_call):
    if not isinstance(tool_call, dict):
        raise ValueError(f"Unsupported Cohere tool call shape: {tool_call!r}")
    function = tool_call.get("function") or {}
    arguments = function.get("arguments", "{}")
    if not isinstance(arguments, str):
        arguments = json.dumps(arguments)
    return {
        "id": str(tool_call.get("id") or f"tool_{len(arguments)}"),
        "type": str(tool_call.get("type") or "function"),
        "function": {
            "name": str(function.get("name") or ""),
            "arguments": arguments,
        },
    }


def clamp_int(value, default, minimum, maximum):
    try:
        number = int(value)
    except (TypeError, ValueError):
        number = default
    return max(minimum, min(maximum, number))


def query_terms(query):
    return [term for term in re.findall(r"[a-z0-9]{3,}", str(query).lower()) if term]


def best_snippet(text, terms, radius=420):
    lowered = text.lower()
    positions = [lowered.find(term) for term in terms if term in lowered]
    position = min([item for item in positions if item >= 0], default=0)
    start = max(0, position - radius)
    end = min(len(text), position + radius)
    return normalize_line(text[start:end])[:900]


def search_extracted_documents(query, top_k=4):
    terms = query_terms(query)
    limit = clamp_int(top_k, 4, 1, 8)
    results = []
    for document in documents:
        text = document.get("text", "")
        if document.get("status") != "extracted" or not text.strip():
            continue
        lowered = text.lower()
        score = sum(lowered.count(term) for term in terms) if terms else 0
        if score <= 0 and terms:
            continue
        results.append({
            "source": document.get("name", ""),
            "url": document.get("url", ""),
            "status": document.get("status", ""),
            "score": score,
            "snippet": best_snippet(text, terms),
        })
    if not results:
        for document in documents:
            text = document.get("text", "")
            if document.get("status") == "extracted" and text.strip():
                results.append({
                    "source": document.get("name", ""),
                    "url": document.get("url", ""),
                    "status": document.get("status", ""),
                    "score": 0,
                    "snippet": best_snippet(text, terms),
                })
            if len(results) >= limit:
                break
    return sorted(results, key=lambda item: item["score"], reverse=True)[:limit]


def get_bid_evidence(evidence_bundle, section, top_k=20):
    evidence = evidence_bundle.get("evidence", {})
    limit = clamp_int(top_k, 20, 1, 40)
    section = str(section).lower()
    if section in {"requirements", "deadlines", "matched_terms"}:
        value = evidence.get(section, [])
        return value[:limit] if isinstance(value, list) else [value]
    if section == "documents":
        return evidence_bundle.get("documents", [])[:limit]
    if section == "opportunity":
        return [evidence_bundle.get("opportunity", {})]
    if section == "profile":
        return [evidence_bundle.get("profile", {})]
    return [{"error": f"unknown evidence section: {section}"}]


def execute_tool_call(tool_call, evidence_bundle):
    normalized = normalize_tool_call(tool_call)
    name = normalized["function"]["name"]
    try:
        arguments = json.loads(normalized["function"]["arguments"] or "{}")
    except json.JSONDecodeError:
        arguments = {}

    if name == "search_extracted_documents":
        result = search_extracted_documents(
            arguments.get("query", ""),
            arguments.get("top_k", 4),
        )
    elif name == "get_bid_evidence":
        result = get_bid_evidence(
            evidence_bundle,
            arguments.get("section", "requirements"),
            arguments.get("top_k", 20),
        )
    else:
        result = [{"error": f"unsupported tool: {name}"}]

    if isinstance(result, dict):
        result_items = [result]
    elif isinstance(result, list):
        result_items = result
    else:
        result_items = [{"value": str(result)}]
    return normalized, arguments, result_items[:10]


def tool_results_message(tool_call_id, result_items):
    content = []
    for index, item in enumerate(result_items):
        content.append({
            "type": "document",
            "document": {
                "id": f"{tool_call_id}:{index}",
                "data": json.dumps(item, ensure_ascii=False),
            },
        })
    return {
        "role": "tool",
        "tool_call_id": tool_call_id,
        "content": content or [{
            "type": "document",
            "document": {"id": f"{tool_call_id}:0", "data": "{}"},
        }],
    }


def synthesize_from_tool_results(api_key, prompt, base_request, tool_result_evidence):
    # Ask Cohere for the final review using plain-text tool evidence.
    # This avoids provider-specific `role: tool` document-block validation while
    # still forcing the first Cohere call to inspect evidence through tools.
    synthesis_prompt = {
        **prompt,
        "tool_results": tool_result_evidence,
        "final_instruction": (
            "The evidence-tool phase is complete. Do not request additional tools. "
            "Return the final strict JSON object now."
        ),
    }
    synthesis_request = {
        **base_request,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are a Canadian public procurement bid-room analyst. "
                    "Use the provided extracted evidence and tool results. Return only valid JSON."
                ),
            },
            {
                "role": "user",
                "content": (
                    "Generate the final JSON object for this bid-room review. "
                    "The first character of your response must be `{`. "
                    "Do not request tools and do not include markdown.\n\n"
                    + json.dumps(synthesis_prompt, ensure_ascii=False)
                ),
            },
        ],
        "response_format": payload.get("cohere", {}).get("response_format", {"type": "json_object"}),
    }
    synthesis_body = cohere_post(api_key, synthesis_request)
    synthesis_content = extract_message_text(synthesis_body.get("message") or {})
    if not synthesis_content:
        raise RuntimeError(
            "Cohere did not produce a final JSON response after tool calls. "
            f"Body snippet: {json.dumps(synthesis_body, ensure_ascii=False)[:900]}"
        )
    return validate_cohere_analysis(parse_model_json(str(synthesis_content)))


def compact_tool_result(item):
    if not isinstance(item, dict):
        return {"value": normalize_line(str(item))[:700]}
    compact = {}
    for key in ("source", "url", "status", "score", "snippet", "text", "section", "error"):
        if key not in item:
            continue
        value = item.get(key)
        if isinstance(value, str):
            compact[key] = normalize_line(value)[:900]
        else:
            compact[key] = value
    if not compact:
        raw = json.dumps(item, ensure_ascii=False)
        compact["value"] = raw[:900]
    return compact


def call_cohere(evidence_bundle):
    api_key = os.environ.get("COHERE_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("COHERE_API_KEY is not available inside the sandbox.")

    prompt = {
        "opportunity": evidence_bundle["opportunity"],
        "profile": evidence_bundle["profile"],
        "documents": evidence_bundle["documents"],
        "deterministic_evidence": evidence_bundle["evidence"],
        "instructions": (
            "Return strict JSON only. Do not use markdown. Do not invent facts. "
            "Use the source names from the evidence when possible."
        ),
        "required_schema": {
            "bid_recommendation": "string: pursue, maybe, or pass with a short reason",
            "fit_score": "integer 0-100",
            "requirements": ["requirement strings with source names when useful"],
            "risks": ["risk strings"],
            "missing_information": ["missing detail strings"],
            "deadlines": ["deadline strings"],
            "questions_to_ask": ["question strings"],
            "next_actions": ["action strings"],
        },
        "output_rule": (
            "Your entire response must be one valid JSON object. The first character "
            "must be { and the last character must be }. Do not include analysis, "
            "markdown, code fences, or text before or after the JSON object."
        ),
        "json_skeleton": {
            "bid_recommendation": "maybe - short reason",
            "fit_score": 0,
            "requirements": [],
            "risks": [],
            "missing_information": [],
            "deadlines": [],
            "questions_to_ask": [],
            "next_actions": [],
        },
    }
    messages = [
        {
            "role": "system",
            "content": (
                "You are a Canadian public procurement bid-room analyst. "
                "You review tender evidence for small businesses. Use the available evidence tools "
                "before producing your final answer. Return only valid JSON in the final answer."
            ),
        },
        {
            "role": "user",
            "content": (
                "First call search_extracted_documents with a query covering requirements, "
                "deadlines, submission, and scope to inspect extracted source text. "
                "Then generate a JSON object for this bid-room review. "
                "The first character of the final response must be `{`. "
                "Do not start the final response with 'We need', 'Here is', or any explanation.\n\n"
                + json.dumps(prompt, ensure_ascii=False)
            ),
        },
    ]
    base_request = {
        "model": payload.get("cohere", {}).get("model", "command-a-plus-05-2026"),
        "max_tokens": int(payload.get("cohere", {}).get("max_tokens", 2400)),
        "temperature": 0.1,
        "p": 0.95,
        "stream": False,
    }
    tool_request = {
        **base_request,
        "messages": messages,
        "tools": payload.get("cohere", {}).get("tools", COHERE_TOOLS),
        "tool_choice": "REQUIRED",
        "strict_tools": True,
    }
    tool_body = cohere_post_with_fallbacks(api_key, tool_request)
    tool_message = tool_body.get("message") or {}
    tool_calls = [normalize_tool_call(item) for item in (tool_message.get("tool_calls") or [])]
    tool_trace = []
    tool_result_evidence = []

    if tool_calls:
        assistant_tool_message = {
            "role": "assistant",
            "tool_calls": tool_calls,
        }
        if tool_message.get("tool_plan"):
            assistant_tool_message["tool_plan"] = str(tool_message.get("tool_plan"))
        messages.append(assistant_tool_message)

        for tool_call in tool_calls[:4]:
            normalized, arguments, result_items = execute_tool_call(tool_call, evidence_bundle)
            compact_results = [compact_tool_result(item) for item in result_items[:3]]
            tool_trace.append({
                "id": normalized["id"],
                "name": normalized["function"]["name"],
                "arguments": arguments,
                "result_count": len(result_items),
                "results": compact_results,
            })
            tool_result_evidence.append({
                "tool": normalized["function"]["name"],
                "arguments": arguments,
                "results": compact_results,
            })
    else:
        tool_trace.append({
            "id": "",
            "name": "none",
            "arguments": {},
            "result_count": 0,
            "error": "Cohere returned no tool calls.",
        })

    if tool_result_evidence:
        return synthesize_from_tool_results(api_key, prompt, base_request, tool_result_evidence), tool_trace

    final_request = {
        **base_request,
        "messages": messages,
        "response_format": payload.get("cohere", {}).get("response_format", {"type": "json_object"}),
    }
    try:
        final_body = cohere_post(api_key, final_request)
    except RuntimeError:
        final_request = {
            **base_request,
            "messages": messages,
            "tools": payload.get("cohere", {}).get("tools", COHERE_TOOLS),
            "tool_choice": "NONE",
        }
        final_body = cohere_post_with_fallbacks(api_key, final_request)

    for _ in range(4):
        message = final_body.get("message") or {}
        content = extract_message_text(message)
        if content:
            return validate_cohere_analysis(parse_model_json(str(content))), tool_trace

        followup_tool_calls = [
            normalize_tool_call(item)
            for item in (message.get("tool_calls") or [])
        ]
        if not followup_tool_calls:
            raise RuntimeError(
                "Cohere returned an empty final message. "
                f"Body snippet: {json.dumps(final_body, ensure_ascii=False)[:900]}"
            )

        assistant_tool_message = {
            "role": "assistant",
            "tool_calls": followup_tool_calls,
        }
        if message.get("tool_plan"):
            assistant_tool_message["tool_plan"] = str(message.get("tool_plan"))
        messages.append(assistant_tool_message)

        for tool_call in followup_tool_calls[:4]:
            normalized, arguments, result_items = execute_tool_call(tool_call, evidence_bundle)
            compact_results = [compact_tool_result(item) for item in result_items[:3]]
            tool_trace.append({
                "id": normalized["id"],
                "name": normalized["function"]["name"],
                "arguments": arguments,
                "result_count": len(result_items),
                "results": compact_results,
            })
            tool_result_evidence.append({
                "tool": normalized["function"]["name"],
                "arguments": arguments,
                "results": compact_results,
            })
            messages.append(tool_results_message(normalized["id"], result_items))

        final_request = {
            **base_request,
            "messages": messages,
            "response_format": payload.get("cohere", {}).get("response_format", {"type": "json_object"}),
        }
        try:
            final_body = cohere_post(api_key, final_request)
        except RuntimeError:
            final_request = {
                **base_request,
                "messages": messages,
                "tools": payload.get("cohere", {}).get("tools", COHERE_TOOLS),
                "tool_choice": "NONE",
            }
            final_body = cohere_post_with_fallbacks(api_key, final_request)

    synthesis_prompt = {
        **prompt,
        "tool_results": tool_result_evidence,
        "final_instruction": (
            "The evidence-tool phase is complete. Do not request additional tools. "
            "Return the final strict JSON object now."
        ),
    }
    synthesis_request = {
        **base_request,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are a Canadian public procurement bid-room analyst. "
                    "Use the provided extracted evidence and tool results. Return only valid JSON."
                ),
            },
            {
                "role": "user",
                "content": (
                    "Generate the final JSON object for this bid-room review. "
                    "The first character of your response must be `{`. "
                    "Do not request tools and do not include markdown.\n\n"
                    + json.dumps(synthesis_prompt, ensure_ascii=False)
                ),
            },
        ],
        "response_format": payload.get("cohere", {}).get("response_format", {"type": "json_object"}),
    }
    synthesis_body = cohere_post(api_key, synthesis_request)
    synthesis_content = extract_message_text(synthesis_body.get("message") or {})
    if not synthesis_content:
        raise RuntimeError(
            "Cohere did not produce a final JSON response after tool calls. "
            f"Body snippet: {json.dumps(synthesis_body, ensure_ascii=False)[:900]}"
        )
    return validate_cohere_analysis(parse_model_json(str(synthesis_content))), tool_trace

documents = []
warnings = []

for inline in payload.get("documents", []):
    text = str(inline.get("text", ""))
    documents.append({
        "name": inline.get("name", "inline-document.txt"),
        "source": inline.get("source", "inline"),
        "url": inline.get("url", ""),
        "bytes": len(text.encode("utf-8")),
        "sha256": sha256_bytes(text.encode("utf-8")),
        "status": "extracted",
        "text": text[:MAX_DOC_CHARS],
        "text_length": len(text),
        "error": "",
    })

for index, attachment in enumerate(payload.get("attachments", [])[: int(payload.get("limits", {}).get("max_attachments", 5))], 1):
    url = attachment.get("url", "")
    name = attachment.get("name") or safe_name(url, f"attachment-{index}")
    record = {
        "name": name,
        "source": attachment.get("kind", "attachment"),
        "url": url,
        "bytes": 0,
        "sha256": "",
        "status": "pending",
        "text": "",
        "text_length": 0,
        "error": "",
    }
    data, headers, error = read_url(url)
    if error:
        record["status"] = "download_failed"
        record["error"] = error
        documents.append(record)
        continue
    record["bytes"] = len(data)
    record["sha256"] = sha256_bytes(data)
    path = download_dir / safe_name(name or url, f"attachment-{index}")
    path.write_bytes(data)
    try:
        text = extract_document_text(path, str(headers.get("Content-Type", "")))
        record["text"] = text[:MAX_DOC_CHARS]
        record["text_length"] = len(text)
        record["status"] = "extracted" if text.strip() else "empty"
    except Exception as exc:
        record["status"] = "extract_failed"
        record["error"] = str(exc)
    documents.append(record)

text_documents = [
    {"name": item["name"], "text": item.get("text", "")}
    for item in documents
    if item.get("text")
]
evidence = extract_evidence(text_documents, payload.get("profile", {}))
document_summaries = [
    {key: item.get(key) for key in ("name", "source", "url", "bytes", "sha256", "status", "text_length", "error")}
    for item in documents
]
evidence_bundle = {
    "opportunity": payload.get("opportunity", {}),
    "profile": payload.get("profile", {}),
    "documents": document_summaries,
    "evidence": {
        "matched_terms": evidence["matched_terms"],
        "requirements": evidence["requirements"],
        "deadlines": evidence["deadlines"],
        "text_for_model": evidence["text_for_model"],
    },
}

cohere_analysis = None
cohere_tool_calls = []
if payload.get("cohere", {}).get("enabled"):
    cohere_analysis, cohere_tool_calls = call_cohere(evidence_bundle)

artifact = {
    "processor": "workspacealberta-e2b-bid-room-v1",
    "processed_at_utc": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
    "opportunity": payload.get("opportunity", {}),
    "profile": payload.get("profile", {}),
    "documents": document_summaries,
    "evidence": {
        "matched_terms": evidence["matched_terms"],
        "requirements": evidence["requirements"],
        "deadlines": evidence["deadlines"],
        "text_characters_sent_to_model": len(evidence["text_for_model"]),
    },
    "cohere_analysis": cohere_analysis,
    "cohere_tool_calls": cohere_tool_calls,
    "warnings": warnings,
}

print(json.dumps(artifact, indent=2, ensure_ascii=False))
"""


@dataclass
class BidRoomSandboxResult:
    """Result returned by a live E2B bid room run."""

    sandbox_id: str
    killed: bool
    artifact: dict[str, Any]
    stdout: str
    stderr: str


def load_local_env() -> None:
    """Load repo-local .env values without printing secrets."""
    for env_path in (ROOT_DIR / ".env", Path.cwd() / ".env"):
        if not env_path.exists():
            continue
        for line in env_path.read_text(encoding="utf-8-sig").splitlines():
            stripped = line.strip()
            if not stripped or stripped.startswith("#") or "=" not in stripped:
                continue
            key, value = stripped.split("=", 1)
            key = key.strip()
            if key.startswith("export "):
                key = key[len("export "):].strip()
            if key and key not in os.environ:
                os.environ[key] = value.strip().strip('"').strip("'")


def has_e2b_api_key() -> bool:
    """Return true when an E2B key is present in the environment or local .env."""
    load_local_env()
    return bool(os.environ.get("E2B_API_KEY", "").strip())


def has_cohere_api_key() -> bool:
    """Return true when the non-prod Cohere key is available."""
    load_local_env()
    return bool(os.environ.get("COHERE_API_KEY", "").strip())


def _field(row: dict[str, Any], *names: str) -> str:
    for name in names:
        value = row.get(name)
        if value:
            return str(value)
    return ""


def split_attachment_urls(*values: str) -> list[str]:
    """Extract and dedupe HTTP(S) URLs from CanadaBuys attachment fields."""
    urls: list[str] = []
    seen = set()
    for value in values:
        for match in re.finditer(r"https?://[^\s,;\"']+", str(value or "")):
            url = match.group(0).strip().rstrip(").]")
            if url and url not in seen:
                seen.add(url)
                urls.append(url)
    return urls


def collect_canadabuys_attachment_urls(
    contract: dict[str, Any],
    *,
    max_attachments: int = MAX_ATTACHMENTS,
) -> list[str]:
    """Collect direct CanadaBuys attachment URLs from known CSV fields."""
    urls = split_attachment_urls(
        _field(contract, "attachment_urls"),
        _field(contract, "attachment-piecesJointes-eng"),
        _field(contract, "attachment-piecesJointes-fra"),
    )
    return urls[:max_attachments]


def _name_from_url(url: str, fallback: str) -> str:
    parsed = urlparse(url)
    name = unquote(Path(parsed.path).name)
    name = re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("._")
    return name[:120] or fallback


def profile_for_bid_room(profile: dict[str, Any] | None, business_context: str = "") -> dict[str, Any]:
    """Normalize saved profile data for sandbox processing."""
    profile = profile or {}
    description = business_context or str(profile.get("description") or "")
    keywords = profile.get("capabilities") or profile.get("keywords") or []
    if not keywords and description:
        keywords = [
            word.lower()
            for word in re.findall(r"[A-Za-z][A-Za-z0-9-]{3,}", description)
        ][:20]
    return {
        "company_name": profile.get("company_name") or "Your Business",
        "location": profile.get("location") or "",
        "description": description or "No business profile was provided.",
        "keywords": [str(item) for item in keywords[:30]],
        "industries": profile.get("industries") or [],
    }


def build_canadabuys_bid_room_payload(
    contract: dict[str, Any],
    profile: dict[str, Any],
    *,
    business_context: str = "",
    max_attachments: int = MAX_ATTACHMENTS,
) -> dict[str, Any]:
    """Build an E2B payload for a CanadaBuys opportunity."""
    reference = _field(contract, "referenceNumber-numeroReference")
    title = _field(contract, "title-titre-eng", "title-titre-fra") or "Untitled CanadaBuys opportunity"
    attachments = [
        {
            "url": url,
            "name": _name_from_url(url, f"canadabuys-attachment-{index}"),
            "kind": "canadabuys_attachment",
        }
        for index, url in enumerate(
            collect_canadabuys_attachment_urls(contract, max_attachments=max_attachments),
            1,
        )
    ]
    description = _field(contract, "tenderDescription-descriptionAppelOffres-eng")
    inline_text = "\n".join(
        part
        for part in (
            f"Title: {title}",
            f"Reference: {reference}",
            f"Solicitation: {_field(contract, 'solicitationNumber-numeroSollicitation')}",
            f"Buyer: {_field(contract, 'contractingEntityName-nomEntitContractante-eng')}",
            f"Closing: {_field(contract, 'tenderClosingDate-appelOffresDateCloture')}",
            f"Category: {_field(contract, 'procurementCategory-categorieApprovisionnement')}",
            f"Region: {_field(contract, 'regionsOfDelivery-regionsLivraison-eng', 'regionsOfOpportunity-regionAppelOffres-eng')}",
            description,
        )
        if part
    )
    return build_process_payload(
        opportunity={
            "source": "CanadaBuys",
            "reference": reference,
            "solicitation": _field(contract, "solicitationNumber-numeroSollicitation"),
            "title": title,
            "buyer": _field(contract, "contractingEntityName-nomEntitContractante-eng"),
            "status": _field(contract, "tenderStatus-appelOffresStatut-eng"),
            "closing": _field(contract, "tenderClosingDate-appelOffresDateCloture"),
            "url": _field(contract, "noticeURL-URLavis-eng"),
        },
        profile=profile_for_bid_room(profile, business_context),
        documents=[{"name": "canadabuys-notice.txt", "text": inline_text, "source": "canadabuys_notice"}],
        attachments=attachments,
    )


def build_apc_bid_room_payload(
    details: dict[str, Any],
    profile: dict[str, Any],
    *,
    business_context: str = "",
    max_attachments: int = MAX_ATTACHMENTS,
) -> dict[str, Any]:
    """Build an E2B payload for an Alberta Purchasing Connection opportunity."""
    opp = details.get("opportunity") or details
    reference = str(opp.get("referenceNumber") or "")
    title = str(opp.get("title") or opp.get("shortTitle") or "Untitled Alberta opportunity")
    text_parts = [
        f"Title: {title}",
        f"Reference: {reference}",
        f"Solicitation: {opp.get('solicitationNumber', '')}",
        f"Buyer: {opp.get('contractingOrganization', '')}",
        f"Closing: {opp.get('closeDateTime', '')}",
        f"Category: {opp.get('categoryCode', '')}",
        f"Region: {', '.join(str(item) for item in (opp.get('regionOfDelivery') or []))}",
        str(opp.get("projectDescription") or ""),
        str(opp.get("additionalRequirements") or ""),
        str(opp.get("submissionDetails") or ""),
        str(opp.get("questionSubmissionDetails") or ""),
        str(opp.get("bidSecurity") or ""),
    ]
    commodity_codes = details.get("commodityCodes") or opp.get("commodityCodes") or []
    if commodity_codes:
        text_parts.append("Commodity codes: " + json.dumps(commodity_codes, ensure_ascii=False))

    attachments = []
    external_link = str(opp.get("externalOriginLink") or "")
    if external_link:
        attachments.append({
            "url": external_link,
            "name": _name_from_url(external_link, "apc-external-page.html"),
            "kind": "apc_external_page",
        })
    return build_process_payload(
        opportunity={
            "source": "Alberta Purchasing Connection",
            "reference": reference,
            "solicitation": str(opp.get("solicitationNumber") or ""),
            "title": title,
            "buyer": str(opp.get("contractingOrganization") or ""),
            "status": str(opp.get("statusCode") or ""),
            "closing": str(opp.get("closeDateTime") or ""),
            "url": external_link,
        },
        profile=profile_for_bid_room(profile, business_context),
        documents=[{
            "name": "alberta-apc-details.txt",
            "text": "\n".join(part for part in text_parts if part),
            "source": "apc_details",
        }],
        attachments=attachments[:max_attachments],
    )


def build_process_payload(
    *,
    opportunity: dict[str, Any],
    profile: dict[str, Any],
    documents: list[dict[str, Any]],
    attachments: list[dict[str, Any]],
    cohere_enabled: bool = True,
) -> dict[str, Any]:
    """Build the normalized sandbox payload."""
    return {
        "opportunity": opportunity,
        "profile": profile,
        "documents": documents,
        "attachments": attachments[:MAX_ATTACHMENTS],
        "limits": {
            "max_attachments": MAX_ATTACHMENTS,
            "max_file_bytes": MAX_FILE_BYTES,
            "max_cohere_chars": MAX_COHERE_CHARS,
        },
        "cohere": {
            "enabled": cohere_enabled,
            "model": COHERE_MODEL,
            "endpoint": COHERE_CHAT_URL,
            "max_tokens": 2400,
            "response_format": COHERE_RESPONSE_FORMAT,
        },
    }


def build_sample_payload(*, cohere_enabled: bool = False) -> dict[str, Any]:
    """Build the sample bid package payload used by the live smoke run."""
    return build_process_payload(
        opportunity={
            "source": "sample",
            "reference": "SAMPLE-BID-ROOM",
            "title": "Sample Structural Steel Tender",
            "buyer": "WorkspaceAlberta",
            "closing": "2026-06-18T14:00:00",
            "url": "",
        },
        profile=profile_for_bid_room({
            "company_name": "Edmonton Steel Works",
            "location": "Edmonton, Alberta",
            "description": (
                "Edmonton steel fabrication shop specializing in structural beams, "
                "railings, custom metalwork, commercial construction, and shop drawings."
            ),
            "capabilities": [
                "steel",
                "structural",
                "fabrication",
                "shop drawings",
                "commercial",
                "edmonton",
            ],
        }),
        documents=[
            {"name": item["name"], "text": item["text"], "source": "sample"}
            for item in SAMPLE_DOCUMENTS
        ],
        attachments=[],
        cohere_enabled=cohere_enabled,
    )


def build_sandbox_command(payload: dict[str, Any]) -> str:
    """Build a Python command that processes a bid package inside E2B."""
    script = SANDBOX_PROCESSOR.replace("__PAYLOAD_JSON__", repr(json.dumps(payload)))
    return "python3 - <<'PY'\n" + script + "\nPY"


def parse_artifact(stdout: str) -> dict[str, Any]:
    """Parse the JSON artifact printed by the sandbox processor."""
    start = stdout.find("{")
    end = stdout.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("Sandbox output did not contain a JSON artifact.")
    return json.loads(stdout[start : end + 1])


def validate_cohere_analysis(data: dict[str, Any]) -> dict[str, Any]:
    """Validate the Cohere JSON shape returned by the sandbox."""
    if not isinstance(data, dict):
        raise ValueError("Cohere analysis must be a JSON object.")
    missing = [field for field in REQUIRED_COHERE_FIELDS if field not in data]
    if missing:
        raise ValueError(f"Cohere analysis missing fields: {', '.join(missing)}")
    if not isinstance(data["bid_recommendation"], str):
        raise ValueError("bid_recommendation must be a string.")
    try:
        data["fit_score"] = int(data["fit_score"])
    except (TypeError, ValueError) as exc:
        raise ValueError("fit_score must be an integer.") from exc
    data["fit_score"] = max(0, min(100, data["fit_score"]))
    for field in REQUIRED_COHERE_FIELDS:
        if field in {"bid_recommendation", "fit_score"}:
            continue
        if isinstance(data[field], str):
            data[field] = [data[field]]
        if not isinstance(data[field], list):
            raise ValueError(f"{field} must be a list.")
    return data


def validate_bid_room_artifact(
    artifact: dict[str, Any],
    *,
    require_cohere: bool = False,
) -> dict[str, Any]:
    """Validate the host-visible bid room artifact."""
    for field in ("processor", "opportunity", "profile", "documents", "evidence"):
        if field not in artifact:
            raise ValueError(f"Bid room artifact missing `{field}`.")
    if not isinstance(artifact["documents"], list):
        raise ValueError("Bid room artifact documents must be a list.")
    if not isinstance(artifact["evidence"], dict):
        raise ValueError("Bid room artifact evidence must be an object.")
    analysis = artifact.get("cohere_analysis")
    if require_cohere and not analysis:
        raise ValueError("Bid room artifact is missing Cohere analysis.")
    if analysis:
        artifact["cohere_analysis"] = validate_cohere_analysis(analysis)
    tool_calls = artifact.get("cohere_tool_calls")
    if require_cohere and not tool_calls:
        raise ValueError("Bid room artifact is missing Cohere tool-call trace.")
    if tool_calls is not None and not isinstance(tool_calls, list):
        raise ValueError("Bid room artifact Cohere tool-call trace must be a list.")
    return artifact


def _result_text(command_result: Any, name: str) -> str:
    value = getattr(command_result, name, "")
    return "" if value is None else str(value)


def _run_e2b_payload(
    payload: dict[str, Any],
    *,
    timeout_seconds: int = 900,
    command_timeout_seconds: int = 420,
    keep_alive: bool = False,
    require_cohere: bool = False,
) -> BidRoomSandboxResult:
    load_local_env()
    if not os.environ.get("E2B_API_KEY", "").strip():
        raise RuntimeError("E2B_API_KEY is not configured.")
    envs = {}
    if payload.get("cohere", {}).get("enabled"):
        cohere_key = os.environ.get("COHERE_API_KEY", "").strip()
        if not cohere_key:
            raise RuntimeError("COHERE_API_KEY is not configured. The prod Cohere key is not used for E2B bid-room processing.")
        envs["COHERE_API_KEY"] = cohere_key

    try:
        from e2b import Sandbox
    except ImportError as exc:
        raise RuntimeError("Install the E2B SDK with `python -m pip install e2b>=2.21.1`.") from exc

    sandbox = Sandbox.create(
        timeout=timeout_seconds,
        envs=envs,
        metadata={
            "project": "workspacealberta",
            "feature": "bid-room-processing",
            "reference": str(payload.get("opportunity", {}).get("reference", ""))[:80],
        },
    )
    sandbox_id = str(
        getattr(sandbox, "sandbox_id", "")
        or getattr(sandbox, "id", "")
        or "unknown"
    )
    artifact: dict[str, Any] | None = None
    stdout = ""
    stderr = ""

    try:
        command = build_sandbox_command(payload)
        command_result = sandbox.commands.run(command, timeout=command_timeout_seconds)
        stdout = _result_text(command_result, "stdout")
        stderr = _result_text(command_result, "stderr")
        artifact = validate_bid_room_artifact(
            parse_artifact(stdout),
            require_cohere=require_cohere,
        )
    finally:
        killed = False
        if not keep_alive:
            sandbox.kill()
            killed = True

    if artifact is None:
        raise RuntimeError("E2B sandbox did not return a bid room artifact.")
    return BidRoomSandboxResult(
        sandbox_id=sandbox_id,
        killed=killed,
        artifact=artifact,
        stdout=stdout,
        stderr=stderr,
    )


def run_live_bid_room_smoke(
    *,
    timeout_seconds: int = 900,
    command_timeout_seconds: int = 120,
    keep_alive: bool = False,
) -> BidRoomSandboxResult:
    """Create a live E2B sandbox, process a sample bid room, and return JSON."""
    return _run_e2b_payload(
        build_sample_payload(cohere_enabled=False),
        timeout_seconds=timeout_seconds,
        command_timeout_seconds=command_timeout_seconds,
        keep_alive=keep_alive,
        require_cohere=False,
    )


def run_live_bid_room_process(
    payload: dict[str, Any],
    *,
    timeout_seconds: int = 900,
    command_timeout_seconds: int = 420,
    keep_alive: bool = False,
) -> BidRoomSandboxResult:
    """Create an E2B sandbox, extract evidence, call Cohere inside it, and return JSON."""
    payload = dict(payload)
    payload["cohere"] = {**payload.get("cohere", {}), "enabled": True}
    return _run_e2b_payload(
        payload,
        timeout_seconds=timeout_seconds,
        command_timeout_seconds=command_timeout_seconds,
        keep_alive=keep_alive,
        require_cohere=True,
    )


def render_bid_room_markdown(result: BidRoomSandboxResult) -> str:
    """Render a bid-room result for MCP users."""
    artifact = result.artifact
    opportunity = artifact.get("opportunity", {})
    evidence = artifact.get("evidence", {})
    analysis = artifact.get("cohere_analysis") or {}
    output = "# Bid Room Analysis\n\n"
    output += f"**Reference:** `{opportunity.get('reference', '')}`\n"
    output += f"**Source:** {opportunity.get('source', '')}\n"
    output += f"**Title:** {opportunity.get('title', '')}\n"
    output += f"**Sandbox:** `{result.sandbox_id}`"
    if result.killed:
        output += " (closed)"
    output += "\n\n"

    if analysis:
        output += "## Cohere Recommendation\n"
        output += f"**Recommendation:** {analysis.get('bid_recommendation', '')}\n"
        output += f"**Fit Score:** {analysis.get('fit_score', '')}\n\n"
        for heading, key in (
            ("Requirements", "requirements"),
            ("Risks", "risks"),
            ("Missing Information", "missing_information"),
            ("Deadlines", "deadlines"),
            ("Questions To Ask", "questions_to_ask"),
            ("Next Actions", "next_actions"),
        ):
            items = analysis.get(key) or []
            output += f"## {heading}\n"
            if items:
                for item in items[:10]:
                    output += f"- {item}\n"
            else:
                output += "- None identified.\n"
            output += "\n"

    output += "## Evidence Processed\n"
    output += f"- **Documents:** {len(artifact.get('documents', []))}\n"
    output += f"- **Matched terms:** {', '.join(evidence.get('matched_terms', [])) or 'None'}\n"
    output += f"- **Requirement-like lines:** {len(evidence.get('requirements', []))}\n"
    output += f"- **Deadline-like lines:** {len(evidence.get('deadlines', []))}\n"
    output += f"- **Characters sent to model:** {evidence.get('text_characters_sent_to_model', 0)}\n\n"
    output += "## Cohere Tool Calls\n"
    tool_calls = artifact.get("cohere_tool_calls") or []
    if tool_calls:
        for tool_call in tool_calls[:8]:
            output += (
                f"- `{tool_call.get('name', '')}` returned "
                f"{tool_call.get('result_count', 0)} result(s)"
            )
            arguments = tool_call.get("arguments") or {}
            if arguments:
                output += f" for `{json.dumps(arguments, ensure_ascii=False)}`"
            output += "\n"
    else:
        output += "- None recorded.\n"
    output += "\n"
    output += "## Document Status\n"
    for document in artifact.get("documents", [])[:12]:
        output += (
            f"- `{document.get('name', '')}`: {document.get('status', '')}, "
            f"{document.get('bytes', 0)} bytes, {document.get('text_length', 0)} text chars"
        )
        if document.get("error"):
            output += f" ({document.get('error')})"
        output += "\n"
    return output.strip()
